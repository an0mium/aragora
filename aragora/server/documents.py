"""
Document parsing and storage for Aragora debates.

Supports PDF, DOCX, TXT, and Markdown files.
Inspired by Heavy3.ai document attachment feature.
"""

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# =============================================================================
# Security: Path Traversal Protection
# =============================================================================

# Pattern for valid document IDs (alphanumeric, hyphens, underscores only)
VALID_DOC_ID_PATTERN = re.compile(r"^[a-zA-Z0-9_-]+$")


def _validate_doc_id(doc_id: str) -> bool:
    """Validate a document ID to prevent path traversal.

    Args:
        doc_id: The document ID to validate

    Returns:
        True if the ID is safe, False otherwise
    """
    if not doc_id:
        return False

    # Must match safe pattern (no path separators, no dots at start)
    if not VALID_DOC_ID_PATTERN.match(doc_id):
        return False

    # Additional checks
    if len(doc_id) > 128:  # Reasonable length limit
        return False

    return True


def _safe_path(storage_dir: Path, doc_id: str) -> Optional[Path]:
    """Construct a safe path for a document, preventing path traversal.

    Args:
        storage_dir: The base storage directory
        doc_id: The document ID

    Returns:
        Safe path if valid, None if the ID is unsafe
    """
    if not _validate_doc_id(doc_id):
        logger.warning(f"Invalid document ID rejected: {doc_id!r}")
        return None

    # Construct path and verify it's within storage_dir
    doc_path = (storage_dir / f"{doc_id}.json").resolve()
    storage_resolved = storage_dir.resolve()

    # Ensure the path is within the storage directory
    try:
        doc_path.relative_to(storage_resolved)
    except ValueError:
        logger.warning(f"Path traversal attempt detected: {doc_id!r}")
        return None

    return doc_path


# Optional PDF support
try:
    from pypdf import PdfReader

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PdfReader = None  # type: ignore[misc,assignment]

# Optional DOCX support
try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None


@dataclass
class ParsedDocument:
    """A parsed document with extracted text."""

    id: str
    filename: str
    content_type: str
    text: str
    page_count: int = 1
    word_count: int = 0
    char_count: int = 0
    created_at: datetime = field(default_factory=datetime.now)
    preview: str = ""

    def __post_init__(self):
        self.word_count = len(self.text.split())
        self.char_count = len(self.text)
        # Generate preview (first 500 chars)
        if not self.preview and self.text:
            self.preview = self.text[:500].strip()
            if len(self.text) > 500:
                self.preview += "..."

    def to_dict(self) -> dict:
        """Convert to JSON-serializable dictionary."""
        return {
            "id": self.id,
            "filename": self.filename,
            "content_type": self.content_type,
            "text": self.text,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "created_at": self.created_at.isoformat(),
            "preview": self.preview,
        }


def generate_doc_id(content: bytes, filename: str) -> str:
    """Generate a unique document ID from content hash."""
    hasher = hashlib.sha256()
    hasher.update(content)
    hasher.update(filename.encode())
    return hasher.hexdigest()[:16]


def parse_pdf(content: bytes, filename: str) -> ParsedDocument:
    """Parse a PDF file and extract text."""
    if not PDF_AVAILABLE:
        raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

    import io

    reader = PdfReader(io.BytesIO(content))

    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    full_text = "\n\n".join(text_parts)

    return ParsedDocument(
        id=generate_doc_id(content, filename),
        filename=filename,
        content_type="application/pdf",
        text=full_text,
        page_count=len(reader.pages),
    )


def parse_docx(content: bytes, filename: str) -> ParsedDocument:
    """Parse a Word document and extract text."""
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install python-docx"
        )

    import io

    doc = DocxDocument(io.BytesIO(content))

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n\n".join(text_parts)

    return ParsedDocument(
        id=generate_doc_id(content, filename),
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        text=full_text,
    )


def parse_text(content: bytes, filename: str) -> ParsedDocument:
    """Parse a plain text or markdown file."""
    # Try UTF-8 first, fall back to latin-1
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")

    content_type = "text/markdown" if filename.lower().endswith(".md") else "text/plain"

    return ParsedDocument(
        id=generate_doc_id(content, filename),
        filename=filename,
        content_type=content_type,
        text=text,
    )


def parse_document(content: bytes, filename: str) -> ParsedDocument:
    """
    Parse a document and extract text based on file extension.

    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    - Text (.txt)
    - Markdown (.md)
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return parse_pdf(content, filename)
    elif filename_lower.endswith(".docx"):
        return parse_docx(content, filename)
    elif filename_lower.endswith((".txt", ".md", ".markdown")):
        return parse_text(content, filename)
    else:
        # Try parsing as plain text
        return parse_text(content, filename)


class DocumentStore:
    """
    Stores and retrieves parsed documents.

    Documents are stored in .nomic/documents/ as JSON.
    """

    def __init__(self, storage_dir: Optional[Path] = None):
        if storage_dir is None:
            storage_dir = Path.cwd() / ".nomic" / "documents"
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)

        # In-memory cache
        self._cache: dict[str, ParsedDocument] = {}

    def add(self, doc: ParsedDocument) -> str:
        """Store a parsed document and return its ID.

        Raises:
            ValueError: If the document ID is invalid (path traversal protection)
        """
        # Validate document ID for path traversal protection
        doc_path = _safe_path(self.storage_dir, doc.id)
        if doc_path is None:
            raise ValueError(f"Invalid document ID: {doc.id!r}")

        # Save to cache
        self._cache[doc.id] = doc

        # Save to disk using validated path
        with open(doc_path, "w") as f:
            json.dump(doc.to_dict(), f, indent=2)

        return doc.id

    def get(self, doc_id: str) -> Optional[ParsedDocument]:
        """Retrieve a document by ID.

        Returns None for invalid document IDs (path traversal protection).
        """
        # Check cache first
        if doc_id in self._cache:
            return self._cache[doc_id]

        # Validate document ID for path traversal protection
        doc_path = _safe_path(self.storage_dir, doc_id)
        if doc_path is None:
            return None

        # Load from disk using validated path
        if not doc_path.exists():
            return None

        with open(doc_path) as f:
            data = json.load(f)

        doc = ParsedDocument(
            id=data["id"],
            filename=data["filename"],
            content_type=data["content_type"],
            text=data["text"],
            page_count=data.get("page_count", 1),
            preview=data.get("preview", ""),
        )

        self._cache[doc_id] = doc
        return doc

    def list_all(self) -> list[dict]:
        """List all stored documents (metadata only)."""
        docs = []
        for doc_path in self.storage_dir.glob("*.json"):
            try:
                with open(doc_path) as f:
                    data = json.load(f)
                docs.append(
                    {
                        "id": data["id"],
                        "filename": data["filename"],
                        "word_count": data.get("word_count", 0),
                        "preview": data.get("preview", "")[:100],
                    }
                )
            except json.JSONDecodeError as e:
                logger.warning(f"Corrupted document file {doc_path}: {e}")
                continue
            except KeyError as e:
                logger.warning(f"Missing required field in {doc_path}: {e}")
                continue
            except (IOError, OSError) as e:
                logger.warning(f"Failed to read document {doc_path}: {e}")
                continue
        return docs

    def delete(self, doc_id: str) -> bool:
        """Delete a document by ID.

        Returns False for invalid document IDs (path traversal protection).
        """
        # Validate document ID (path traversal protection)
        doc_path = _safe_path(self.storage_dir, doc_id)
        if doc_path is None:
            return False

        # Remove from cache
        self._cache.pop(doc_id, None)

        # Remove from disk using validated path
        if doc_path.exists():
            doc_path.unlink()
            return True
        return False

    def get_context_for_debate(self, doc_ids: list[str]) -> str:
        """
        Get combined text from multiple documents for debate context.

        Returns formatted text suitable for inclusion in agent prompts.
        """
        if not doc_ids:
            return ""

        parts = []
        for doc_id in doc_ids:
            doc = self.get(doc_id)
            if doc:
                parts.append(f"=== Document: {doc.filename} ===\n{doc.text}")

        return "\n\n".join(parts)


# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


def get_supported_formats() -> dict:
    """Get information about supported document formats."""
    return {
        "formats": [
            {"ext": ".pdf", "mime": "application/pdf", "available": PDF_AVAILABLE},
            {
                "ext": ".docx",
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "available": DOCX_AVAILABLE,
            },
            {"ext": ".txt", "mime": "text/plain", "available": True},
            {"ext": ".md", "mime": "text/markdown", "available": True},
        ],
        "max_size_mb": 10,
    }
