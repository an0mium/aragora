"""
Unstructured.io adapter for document parsing.

Provides unified document parsing for 25+ formats including:
- PDF (with OCR support)
- Microsoft Office (DOCX, XLSX, PPTX)
- Email (EML, MSG)
- Images (PNG, JPG with OCR)
- HTML, XML, JSON, CSV
- Plain text, Markdown, RST

Falls back to native parsers when unstructured is not available.

Usage:
    from aragora.documents.ingestion.unstructured_adapter import (
        UnstructuredParser,
        parse_document,
        get_supported_formats,
    )

    parser = UnstructuredParser()
    result = parser.parse(content, filename="report.pdf")

    # Or use the convenience function
    result = parse_document(content, "report.pdf")
"""

from __future__ import annotations

import io
import logging
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from aragora.documents.models import (
    ChunkType,
    DocumentStatus,
    IngestedDocument,
)

logger = logging.getLogger(__name__)


# Check for unstructured library
try:
    from unstructured.partition.auto import partition  # noqa: F401
    from unstructured.partition.pdf import partition_pdf  # noqa: F401
    from unstructured.partition.docx import partition_docx  # noqa: F401
    from unstructured.partition.html import partition_html  # noqa: F401
    from unstructured.partition.text import partition_text  # noqa: F401
    from unstructured.partition.md import partition_md  # noqa: F401
    from unstructured.partition.json import partition_json  # noqa: F401
    from unstructured.partition.csv import partition_csv  # noqa: F401
    from unstructured.partition.xml import partition_xml  # noqa: F401
    from unstructured.partition.pptx import partition_pptx  # noqa: F401
    from unstructured.partition.xlsx import partition_xlsx  # noqa: F401
    from unstructured.partition.email import partition_email  # noqa: F401
    from unstructured.partition.image import partition_image  # noqa: F401
    from unstructured.documents.elements import (  # noqa: F401
        Element,
        Title,
        NarrativeText,
        ListItem,
        Table,
        Image,
        Header,
        Footer,
        PageBreak,
        Formula,
    )

    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    logger.info("unstructured not available - using fallback parsers")


# Check for native fallback parsers
try:
    from pypdf import PdfReader  # noqa: F401

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument  # noqa: F401

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Supported formats and their configurations
SUPPORTED_FORMATS = {
    # Documents
    ".pdf": {"mime": "application/pdf", "parser": "pdf", "ocr": True},
    ".docx": {
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "parser": "docx",
    },
    ".doc": {"mime": "application/msword", "parser": "doc"},
    ".pptx": {
        "mime": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "parser": "pptx",
    },
    ".ppt": {"mime": "application/vnd.ms-powerpoint", "parser": "ppt"},
    ".xlsx": {
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "parser": "xlsx",
    },
    ".xls": {"mime": "application/vnd.ms-excel", "parser": "xls"},
    # Text
    ".txt": {"mime": "text/plain", "parser": "text"},
    ".md": {"mime": "text/markdown", "parser": "md"},
    ".markdown": {"mime": "text/markdown", "parser": "md"},
    ".rst": {"mime": "text/x-rst", "parser": "rst"},
    # Web
    ".html": {"mime": "text/html", "parser": "html"},
    ".htm": {"mime": "text/html", "parser": "html"},
    ".xml": {"mime": "application/xml", "parser": "xml"},
    # Data
    ".json": {"mime": "application/json", "parser": "json"},
    ".csv": {"mime": "text/csv", "parser": "csv"},
    # Email
    ".eml": {"mime": "message/rfc822", "parser": "email"},
    ".msg": {"mime": "application/vnd.ms-outlook", "parser": "msg"},
    # Images (OCR)
    ".png": {"mime": "image/png", "parser": "image", "ocr": True},
    ".jpg": {"mime": "image/jpeg", "parser": "image", "ocr": True},
    ".jpeg": {"mime": "image/jpeg", "parser": "image", "ocr": True},
    ".tiff": {"mime": "image/tiff", "parser": "image", "ocr": True},
    ".bmp": {"mime": "image/bmp", "parser": "image", "ocr": True},
    # Code
    ".py": {"mime": "text/x-python", "parser": "text"},
    ".js": {"mime": "application/javascript", "parser": "text"},
    ".ts": {"mime": "application/typescript", "parser": "text"},
    ".java": {"mime": "text/x-java-source", "parser": "text"},
    ".cpp": {"mime": "text/x-c++src", "parser": "text"},
    ".c": {"mime": "text/x-csrc", "parser": "text"},
    ".go": {"mime": "text/x-go", "parser": "text"},
    ".rs": {"mime": "text/x-rust", "parser": "text"},
    ".rb": {"mime": "application/x-ruby", "parser": "text"},
}


@dataclass
class ParsedElement:
    """A single parsed element from a document."""

    type: ChunkType
    text: str
    page_number: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParseResult:
    """Result of parsing a document."""

    elements: list[ParsedElement]
    full_text: str
    page_count: int
    headings: list[str]
    tables_count: int
    images_count: int
    parser_used: str
    parse_duration_ms: int
    metadata: dict[str, Any] = field(default_factory=dict)


class UnstructuredParser:
    """
    Document parser using Unstructured.io.

    Provides consistent parsing across 25+ document formats with
    fallback to native parsers when unstructured is not available.
    """

    def __init__(
        self,
        enable_ocr: bool = True,
        languages: list[str] | None = None,
        extract_tables: bool = True,
        extract_images: bool = True,
    ):
        """
        Initialize parser.

        Args:
            enable_ocr: Enable OCR for images and scanned PDFs
            languages: Languages for OCR (default: ["eng"])
            extract_tables: Extract table structures
            extract_images: Extract image captions/OCR
        """
        self.enable_ocr = enable_ocr
        self.languages = languages or ["eng"]
        self.extract_tables = extract_tables
        self.extract_images = extract_images

    def parse(
        self,
        content: bytes,
        filename: str,
        content_type: Optional[str] = None,
    ) -> ParseResult:
        """
        Parse a document and extract structured content.

        Args:
            content: Raw file content
            filename: Original filename (used for format detection)
            content_type: Optional MIME type override

        Returns:
            ParseResult with extracted elements
        """
        start_time = time.monotonic()

        # Determine file extension
        ext = Path(filename).suffix.lower()

        # Get format configuration
        format_config = SUPPORTED_FORMATS.get(ext, {"parser": "text"})

        if UNSTRUCTURED_AVAILABLE:
            result = self._parse_with_unstructured(content, filename, format_config)
        else:
            result = self._parse_with_fallback(content, filename, ext)

        # Calculate duration
        duration_ms = int((time.monotonic() - start_time) * 1000)
        result.parse_duration_ms = duration_ms

        return result

    def _parse_with_unstructured(
        self,
        content: bytes,
        filename: str,
        format_config: dict,
    ) -> ParseResult:
        """Parse using unstructured library."""
        # Create file-like object
        file_obj = io.BytesIO(content)

        # Configure partitioning options
        kwargs: dict[str, Any] = {
            "file": file_obj,
            "metadata_filename": filename,
        }

        # Add OCR settings for supported formats
        if format_config.get("ocr") and self.enable_ocr:
            kwargs["languages"] = self.languages
            kwargs["strategy"] = "hi_res"  # High-resolution OCR

        # Partition the document
        try:
            elements = partition(**kwargs)
        except Exception as e:
            logger.warning(f"unstructured partition failed: {e}, trying fallback")
            ext = Path(filename).suffix.lower()
            return self._parse_with_fallback(content, filename, ext)

        # Convert elements to our format
        parsed_elements = []
        headings = []
        tables_count = 0
        images_count = 0
        page_numbers: set[int] = set()

        for elem in elements:
            chunk_type = self._element_to_chunk_type(elem)
            text = str(elem)

            # Track page numbers
            page_num = getattr(elem.metadata, "page_number", 0) or 0
            if page_num:
                page_numbers.add(page_num)

            # Track headings
            if chunk_type == ChunkType.HEADING:
                headings.append(text)

            # Track tables and images
            if chunk_type == ChunkType.TABLE:
                tables_count += 1
            elif chunk_type == ChunkType.IMAGE:
                images_count += 1

            # Build metadata
            elem_metadata = {}
            if hasattr(elem, "metadata"):
                for attr in ["filename", "page_number", "coordinates", "languages"]:
                    if hasattr(elem.metadata, attr):
                        value = getattr(elem.metadata, attr)
                        if value is not None:
                            elem_metadata[attr] = value

            parsed_elements.append(
                ParsedElement(
                    type=chunk_type,
                    text=text,
                    page_number=page_num,
                    metadata=elem_metadata,
                )
            )

        # Combine text
        full_text = "\n\n".join(elem.text for elem in parsed_elements if elem.text.strip())

        return ParseResult(
            elements=parsed_elements,
            full_text=full_text,
            page_count=max(page_numbers) if page_numbers else 1,
            headings=headings,
            tables_count=tables_count,
            images_count=images_count,
            parser_used="unstructured",
            parse_duration_ms=0,  # Set by caller
        )

    def _element_to_chunk_type(self, elem) -> ChunkType:
        """Convert unstructured element to ChunkType."""
        if not UNSTRUCTURED_AVAILABLE:
            return ChunkType.TEXT

        if isinstance(elem, Title):
            return ChunkType.HEADING
        elif isinstance(elem, (Header, Footer)):
            return ChunkType.METADATA
        elif isinstance(elem, Table):
            return ChunkType.TABLE
        elif isinstance(elem, Image):
            return ChunkType.IMAGE
        elif isinstance(elem, ListItem):
            return ChunkType.LIST
        elif isinstance(elem, Formula):
            return ChunkType.FORMULA
        elif isinstance(elem, NarrativeText):
            return ChunkType.TEXT
        else:
            return ChunkType.TEXT

    def _parse_with_fallback(
        self,
        content: bytes,
        filename: str,
        ext: str,
    ) -> ParseResult:
        """Parse using native fallback parsers."""
        if ext == ".pdf" and PYPDF_AVAILABLE:
            return self._parse_pdf_fallback(content, filename)
        elif ext == ".docx" and DOCX_AVAILABLE:
            return self._parse_docx_fallback(content, filename)
        else:
            return self._parse_text_fallback(content, filename)

    def _parse_pdf_fallback(self, content: bytes, filename: str) -> ParseResult:
        """Parse PDF using pypdf."""
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))

        elements = []
        headings = []
        page_texts = []

        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            page_texts.append(text)

            # Simple heading extraction (lines that look like titles)
            for line in text.split("\n"):
                line = line.strip()
                if line and len(line) < 100 and line.isupper():
                    headings.append(line)
                    elements.append(
                        ParsedElement(
                            type=ChunkType.HEADING,
                            text=line,
                            page_number=i,
                        )
                    )
                elif line:
                    elements.append(
                        ParsedElement(
                            type=ChunkType.TEXT,
                            text=line,
                            page_number=i,
                        )
                    )

        return ParseResult(
            elements=elements,
            full_text="\n\n".join(page_texts),
            page_count=len(reader.pages),
            headings=headings[:50],  # Limit headings
            tables_count=0,  # pypdf doesn't extract tables
            images_count=0,
            parser_used="pypdf",
            parse_duration_ms=0,
        )

    def _parse_docx_fallback(self, content: bytes, filename: str) -> ParseResult:
        """Parse DOCX using python-docx."""
        from docx import Document

        doc = Document(io.BytesIO(content))

        elements = []
        headings = []
        tables_count = len(doc.tables)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if it's a heading
            if para.style and "Heading" in para.style.name:
                headings.append(text)
                elements.append(ParsedElement(type=ChunkType.HEADING, text=text))
            else:
                elements.append(ParsedElement(type=ChunkType.TEXT, text=text))

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                rows.append(row_text)
            table_text = "\n".join(rows)
            elements.append(ParsedElement(type=ChunkType.TABLE, text=table_text))

        full_text = "\n\n".join(e.text for e in elements if e.text)

        return ParseResult(
            elements=elements,
            full_text=full_text,
            page_count=1,  # python-docx doesn't provide page count
            headings=headings,
            tables_count=tables_count,
            images_count=0,
            parser_used="python-docx",
            parse_duration_ms=0,
        )

    def _parse_text_fallback(self, content: bytes, filename: str) -> ParseResult:
        """Parse text files."""
        # Try UTF-8 first, fall back to latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        # Simple heading extraction for markdown
        headings = []
        ext = Path(filename).suffix.lower()

        if ext in (".md", ".markdown"):
            import re

            for match in re.finditer(r"^#{1,6}\s+(.+)$", text, re.MULTILINE):
                headings.append(match.group(1).strip())

        elements = [ParsedElement(type=ChunkType.TEXT, text=text)]

        return ParseResult(
            elements=elements,
            full_text=text,
            page_count=1,
            headings=headings,
            tables_count=0,
            images_count=0,
            parser_used="native-text",
            parse_duration_ms=0,
        )

    def parse_to_document(
        self,
        content: bytes,
        filename: str,
        workspace_id: str = "",
        uploaded_by: str = "",
        tags: Optional[list[str]] = None,
    ) -> IngestedDocument:
        """
        Parse content and return an IngestedDocument.

        Args:
            content: Raw file content
            filename: Original filename
            workspace_id: Workspace/organization ID
            uploaded_by: User ID who uploaded
            tags: Optional tags

        Returns:
            IngestedDocument ready for further processing
        """
        # Parse the document
        result = self.parse(content, filename)

        # Determine content type
        ext = Path(filename).suffix.lower()
        format_config = SUPPORTED_FORMATS.get(ext, {})
        content_type = format_config.get("mime", "application/octet-stream")

        # Create IngestedDocument
        doc = IngestedDocument(
            filename=filename,
            content_type=content_type,
            file_size=len(content),
            workspace_id=workspace_id,
            uploaded_by=uploaded_by,
            status=DocumentStatus.PROCESSING,
            page_count=result.page_count,
            word_count=len(result.full_text.split()),
            char_count=len(result.full_text),
            parser_used=result.parser_used,
            parse_duration_ms=result.parse_duration_ms,
            text=result.full_text,
            headings=result.headings,
            tables_count=result.tables_count,
            images_count=result.images_count,
            tags=tags or [],
        )

        return doc


# Convenience functions
def parse_document(
    content: bytes,
    filename: str,
    enable_ocr: bool = True,
) -> ParseResult:
    """
    Parse a document using the best available method.

    Args:
        content: Raw file content
        filename: Original filename
        enable_ocr: Enable OCR for images/scanned PDFs

    Returns:
        ParseResult with extracted content
    """
    parser = UnstructuredParser(enable_ocr=enable_ocr)
    return parser.parse(content, filename)


def get_supported_formats() -> dict[str, Any]:
    """Get information about supported document formats."""
    formats = []
    for ext, config in SUPPORTED_FORMATS.items():
        formats.append(
            {
                "extension": ext,
                "mime": config.get("mime", ""),
                "ocr_support": config.get("ocr", False),
                "available": _check_format_available(ext),
            }
        )

    return {
        "formats": formats,
        "total_formats": len(formats),
        "unstructured_available": UNSTRUCTURED_AVAILABLE,
        "ocr_available": UNSTRUCTURED_AVAILABLE,  # OCR requires unstructured
        "max_size_mb": 100,  # With streaming support
    }


def _check_format_available(ext: str) -> bool:
    """Check if a format is available for parsing."""
    if UNSTRUCTURED_AVAILABLE:
        return True

    # Check fallback availability
    if ext == ".pdf":
        return PYPDF_AVAILABLE
    elif ext == ".docx":
        return DOCX_AVAILABLE
    elif ext in (".txt", ".md", ".markdown", ".py", ".js", ".ts", ".json"):
        return True

    return False


__all__ = [
    "UnstructuredParser",
    "ParseResult",
    "ParsedElement",
    "parse_document",
    "get_supported_formats",
    "SUPPORTED_FORMATS",
    "UNSTRUCTURED_AVAILABLE",
]
